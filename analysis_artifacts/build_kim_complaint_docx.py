# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT_DIR = Path(r"C:\Users\baek_\newmso\legal_drafts")
DOCX_PATH = OUT_DIR / "고소장_김대운_급여편취_초안_2026-05-11.docx"
MD_PATH = OUT_DIR / "고소장_김대운_급여편취_초안_2026-05-11.md"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, size=9.5, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(14 if level == 1 else 11.5)
    run.font.color.rgb = RGBColor(40, 40, 40)
    return p


def add_body(doc, text, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(10)
    return p


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.first_line_indent = Cm(-0.4)
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        run.font.name = "맑은 고딕"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        run.font.size = Pt(10)


def add_table(doc, rows, widths=None, header=True):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            if widths:
                cell.width = Cm(widths[c_idx])
            is_header = header and r_idx == 0
            set_cell_text(cell, text, bold=is_header, size=9 if len(text) > 36 else 9.5)
            if is_header:
                set_cell_shading(cell, "EDEDED")
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif c_idx in (0, 1):
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


complaint_md = """# 고소장

## 1. 고소인

- 회사명: [주식회사/개인사업자명]
- 대표자: [대표자 성명]
- 사업자등록번호: [사업자등록번호]
- 주소: [본점/사업장 주소]
- 담당자 및 연락처: [성명 / 직책 / 전화번호]

## 2. 피고소인

- 성명: 김대운
- 직위: 사원(영업사원)
- 재직기간: 2026. 3. 3.부터 2026. 4. 20.까지
- 근무시간: 08:30부터 18:00까지
- 주소/연락처/주민등록번호: [인사기록에 기재된 사항 입력]

## 3. 죄명

주위적으로 사기, 예비적으로 법인카드·차량·유류비 등 사적 사용 부분에 관하여 업무상배임 또는 업무상횡령 혐의가 있는지 수사를 요청합니다.

## 4. 고소취지

피고소인은 영업사원으로 근무하면서 실제 근무시간 중 상당 시간을 영업활동에 사용하지 않았음에도, 회사에는 정상적으로 거래처·관공서·병의원 등을 방문하여 영업업무를 수행한 것처럼 차량운행일지 및 보고를 하였습니다. 고소인은 그 허위 보고를 신뢰하여 급여 및 업무 관련 비용을 지급하였거나 지급대상으로 처리하였으므로, 피고소인을 사기 및 관련 범죄로 철저히 수사하여 처벌하여 주시기 바랍니다.

## 5. 범죄사실

피고소인은 2026. 3. 3.경 고소인 회사에 영업사원으로 입사하여 2026. 4. 20.까지 근무하였고, 근무시간 중 거래처·관공서·병의원 방문 등 영업활동을 수행하며 회사 차량 운행일지 및 회사 메신저 등으로 업무 진행상황을 정확히 보고할 의무가 있었습니다.

그럼에도 피고소인은 2026. 4. 13.부터 2026. 4. 16.까지 아래와 같이 차량운행일지에는 정상 영업활동을 한 것처럼 방문지와 업무내역을 기재하였으나, 차량관리 앱 GPS 운행기록 및 차량 블랙박스 캡처를 대조하면 장시간 거주지 인근 또는 특정 장소에 정차해 있었고, 일부 시간에는 개인 용무를 본 정황이 확인됩니다.

1. 2026. 4. 13. 피고소인은 차량운행일지에 ‘무안 군청, 보건소’ 방문으로 기재하였으나, 차량관리 앱상 실제 운행은 09:19~09:35 4.0km 및 16:10~16:20 2.0km에 그쳤고, 블랙박스에는 09:35~16:10경 피고소인의 거주지인 목포시 연산주공3차 인근에 차량이 주차되어 있던 정황이 확인됩니다.
2. 2026. 4. 14. 피고소인은 차량운행일지에 ‘해제, 망운 병의원’ 방문으로 기재하였으나, 차량관리 앱과 블랙박스에 의하면 무안군 소재 식당에서 점심식사를 한 뒤 특정 장소로 이동하여 12:06~16:13경까지 차량 이동이 없었습니다. 위 점심 식대는 법인카드로 결제된 것으로 확인됩니다.
3. 2026. 4. 15. 피고소인은 차량운행일지에 ‘일로 개인병원, 요양원’ 방문으로 기재하였으나, 차량관리 앱과 블랙박스에는 12:28~16:01경까지 같은 장소에 정차되어 있던 정황이 확인됩니다.
4. 2026. 4. 16. 피고소인은 차량운행일지에 ‘완도 군청, 보건소’ 방문으로 기재하였고 실제 완도 방문은 있었던 것으로 보이나, 회사로 바로 복귀하지 않고 근무시간 중 상동 평화주방, 하이마트 등 개인 용무로 보이는 장소를 방문한 정황이 확인됩니다.

또한 회사는 피고소인에게 영업 진행상황을 성실하게 보고할 것을 회사 메신저로 지시하였으나, 피고소인은 허위로 답변하거나 정당한 이유 없이 보고를 누락·거부하였습니다.

피고소인의 위 행위는 단순한 업무태만을 넘어, 회사가 정상 근무 및 정상 영업활동을 한 것으로 오인하도록 허위 운행일지와 보고를 제출한 행위입니다. 고소인은 그 허위 보고를 기초로 피고소인에게 급여 및 업무 관련 비용을 지급하거나 지급대상으로 처리하였고, 이로써 고소인 회사에 재산상 손해가 발생하였습니다.

## 6. 피해내용 및 피해액

현재 확인된 피해는 2026. 4. 13.부터 2026. 4. 16.까지의 허위 또는 부실 근무 시간 상당 급여, 법인카드 식대, 회사 차량 운행 관련 유류비·차량유지비 등입니다. 정확한 피해액은 급여대장, 급여 이체내역, 법인카드 사용내역, 차량 유류비 내역을 제출하여 수사 과정에서 특정하겠습니다.

- 2026. 4월분 지급 급여: [금액]원
- 허위 또는 부실 근무 시간 상당 급여: [산정금액]원
- 법인카드 식대 및 업무비용: [금액]원
- 회사 차량·유류비 등 관련 비용: [금액]원
- 합계: [금액]원

## 7. 증거자료

1. 증 제1호 차량운행일지 사본: 피고소인이 직접 작성한 방문지 및 업무내역 기재 내용
2. 증 제2호 차량관리 앱 운행기록: GPS 기반 이동시간, 주행거리, 장시간 미운행 정황
3. 증 제3호 차량 블랙박스 캡처: 실제 차량 위치 및 장시간 정차 정황
4. 증 제4호 회사 메신저 캡처: 업무보고 지시, 보고 누락·거부 및 허위답변 정황
5. 추가 제출 예정: 근로계약서, 취업규칙 또는 영업사원 업무지시서, 급여대장, 급여 이체내역, 법인카드 사용내역, 유류비 내역, 원본 블랙박스 영상 및 차량관리 앱 원본 로그

## 8. 수사 요청사항

1. 피고소인이 위 날짜들에 실제 방문했다고 기재한 무안 군청·보건소, 해제·망운 병의원, 일로 개인병원·요양원, 완도 군청·보건소 등에 실제 방문·영업활동을 했는지 확인하여 주십시오.
2. 차량관리 앱 원본 로그, 블랙박스 원본 파일, 차량 GPS 기록의 생성시각 및 위치정보를 확인하여 주십시오.
3. 법인카드 사용내역 및 급여 지급내역을 확인하여 피고소인이 허위 보고를 통해 취득한 급여 및 비용 상당액을 특정하여 주십시오.
4. 피고소인이 같은 방식으로 2026. 3. 3.부터 2026. 4. 20.까지 다른 근무일에도 허위 보고를 반복하였는지 수사하여 주십시오.
5. 법인카드 식대, 회사 차량, 유류비 등 업무용 자산을 사적 목적으로 사용한 부분에 대하여 업무상배임 또는 업무상횡령 성립 여부도 함께 검토하여 주십시오.

## 9. 첨부서류

- 김대운 증거 설명서.pdf 1부
- 차량운행일지 사본 1부
- 차량관리 앱 운행기록 캡처 1부
- 차량 블랙박스 캡처 및 원본 영상 저장매체 1부
- 회사 메신저 캡처 1부
- 근로계약서 및 업무지시 관련 자료 1부
- 급여대장 및 급여 이체내역 1부
- 법인카드 사용내역 및 유류비 내역 1부

2026. 5. [  ].

고소인: [회사명/대표자명] (서명 또는 인)

[관할 경찰서장] 귀중
"""


def build_docx():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = doc.styles
    styles["Normal"].font.name = "맑은 고딕"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    styles["Normal"].font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(16)
    run = title.add_run("고 소 장")
    run.bold = True
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(20)

    add_heading(doc, "1. 고소인 및 피고소인", level=1)
    add_table(
        doc,
        [
            ["구분", "내용"],
            ["고소인", "회사명: [주식회사/개인사업자명]\n대표자: [대표자 성명]\n사업자등록번호: [사업자등록번호]\n주소: [본점/사업장 주소]\n담당자 및 연락처: [성명 / 직책 / 전화번호]"],
            ["피고소인", "성명: 김대운\n직위: 사원(영업사원)\n재직기간: 2026. 3. 3.부터 2026. 4. 20.까지\n근무시간: 08:30부터 18:00까지\n주소/연락처/주민등록번호: [인사기록에 기재된 사항 입력]"],
        ],
        widths=[3.0, 13.0],
    )

    add_heading(doc, "2. 죄명", level=1)
    add_body(doc, "주위적으로 사기, 예비적으로 법인카드·차량·유류비 등 사적 사용 부분에 관하여 업무상배임 또는 업무상횡령 혐의가 있는지 수사를 요청합니다.")

    add_heading(doc, "3. 고소취지", level=1)
    add_body(
        doc,
        "피고소인은 영업사원으로 근무하면서 실제 근무시간 중 상당 시간을 영업활동에 사용하지 않았음에도, 회사에는 정상적으로 거래처·관공서·병의원 등을 방문하여 영업업무를 수행한 것처럼 차량운행일지 및 보고를 하였습니다. 고소인은 그 허위 보고를 신뢰하여 급여 및 업무 관련 비용을 지급하였거나 지급대상으로 처리하였으므로, 피고소인을 사기 및 관련 범죄로 철저히 수사하여 처벌하여 주시기 바랍니다.",
    )

    add_heading(doc, "4. 범죄사실", level=1)
    add_body(
        doc,
        "피고소인은 2026. 3. 3.경 고소인 회사에 영업사원으로 입사하여 2026. 4. 20.까지 근무하였고, 근무시간 중 거래처·관공서·병의원 방문 등 영업활동을 수행하며 회사 차량 운행일지 및 회사 메신저 등으로 업무 진행상황을 정확히 보고할 의무가 있었습니다.",
    )
    add_body(
        doc,
        "그럼에도 피고소인은 2026. 4. 13.부터 2026. 4. 16.까지 아래와 같이 차량운행일지에는 정상 영업활동을 한 것처럼 방문지와 업무내역을 기재하였으나, 차량관리 앱 GPS 운행기록 및 차량 블랙박스 캡처를 대조하면 장시간 거주지 인근 또는 특정 장소에 정차해 있었고, 일부 시간에는 개인 용무를 본 정황이 확인됩니다.",
    )

    add_table(
        doc,
        [
            ["일자", "운행일지 기재", "앱/블랙박스상 확인내용", "고소장에 반영할 의미"],
            ["2026. 4. 13.", "무안 군청, 보건소", "실제 운행은 09:19~09:35 4.0km 및 16:10~16:20 2.0km. 09:35~16:10경 목포시 연산주공3차 인근 주차 정황.", "정상 영업활동 보고와 실제 장시간 거주지 인근 체류 정황의 불일치"],
            ["2026. 4. 14.", "해제, 망운 병의원", "09:35~10:48 37.0km, 11:25~12:06 35.0km, 16:13~16:24 3.0km. 식사 후 12:06~16:13경까지 차량 이동 없음. 점심 식대 법인카드 이용 정황.", "허위·부실 근무 및 업무비용 사적 사용 여부"],
            ["2026. 4. 15.", "일로 개인병원, 요양원", "10:28~11:14 16.0km, 12:00~12:27 17.0km, 16:19~16:32 3.0km, 17:01~17:08 2.0km. 12:28~16:01경 같은 장소 정차.", "방문지 보고와 장시간 정차 정황 불일치"],
            ["2026. 4. 16.", "완도 군청, 보건소", "완도 방문은 있었던 것으로 보이나, 회사로 곧바로 복귀하지 않고 근무시간 중 상동 평화주방, 하이마트 등 개인 용무 장소 방문 정황.", "근무시간 중 사적 용무 및 보고의 충실성 문제"],
        ],
        widths=[2.4, 3.1, 6.8, 3.7],
    )

    add_body(
        doc,
        "또한 회사는 피고소인에게 영업 진행상황을 성실하게 보고할 것을 회사 메신저로 지시하였으나, 피고소인은 허위로 답변하거나 정당한 이유 없이 보고를 누락·거부하였습니다.",
    )
    add_body(
        doc,
        "피고소인의 위 행위는 단순한 업무태만을 넘어, 회사가 정상 근무 및 정상 영업활동을 한 것으로 오인하도록 허위 운행일지와 보고를 제출한 행위입니다. 고소인은 그 허위 보고를 기초로 피고소인에게 급여 및 업무 관련 비용을 지급하거나 지급대상으로 처리하였고, 이로써 고소인 회사에 재산상 손해가 발생하였습니다.",
    )

    add_heading(doc, "5. 피해내용 및 피해액", level=1)
    add_body(
        doc,
        "현재 확인된 피해는 2026. 4. 13.부터 2026. 4. 16.까지의 허위 또는 부실 근무 시간 상당 급여, 법인카드 식대, 회사 차량 운행 관련 유류비·차량유지비 등입니다. 정확한 피해액은 급여대장, 급여 이체내역, 법인카드 사용내역, 차량 유류비 내역을 제출하여 수사 과정에서 특정하겠습니다.",
    )
    add_table(
        doc,
        [
            ["항목", "금액", "비고"],
            ["2026. 4월분 지급 급여", "[금액]원", "급여대장 및 이체내역 첨부"],
            ["허위 또는 부실 근무 시간 상당 급여", "[산정금액]원", "시간급 환산 또는 월급 일할·시간할 기준으로 산정"],
            ["법인카드 식대 및 업무비용", "[금액]원", "2026. 4. 14. 식대 등"],
            ["회사 차량·유류비 등 관련 비용", "[금액]원", "차량관리 앱, 유류비 내역 기준"],
            ["합계", "[금액]원", "수사 과정에서 최종 특정"],
        ],
        widths=[5.0, 4.0, 7.0],
    )

    add_heading(doc, "6. 증거자료", level=1)
    add_table(
        doc,
        [
            ["증거번호", "증거명", "입증취지"],
            ["증 제1호", "차량운행일지 사본", "피고소인이 직접 작성한 방문지 및 업무내역 기재 내용"],
            ["증 제2호", "차량관리 앱 운행기록", "GPS 기반 이동시간, 주행거리, 장시간 미운행 정황"],
            ["증 제3호", "차량 블랙박스 캡처", "실제 차량 위치 및 장시간 정차 정황"],
            ["증 제4호", "회사 메신저 캡처", "업무보고 지시, 보고 누락·거부 및 허위답변 정황"],
            ["추가", "근로계약서, 급여대장, 이체내역, 법인카드 내역, 유류비 내역, 원본 블랙박스 영상 및 차량관리 앱 원본 로그", "기망에 따른 급여·비용 지급 및 피해액 특정"],
        ],
        widths=[2.4, 6.0, 7.6],
    )

    add_heading(doc, "7. 수사 요청사항", level=1)
    add_numbered(
        doc,
        [
            "1. 피고소인이 위 날짜들에 실제 방문했다고 기재한 무안 군청·보건소, 해제·망운 병의원, 일로 개인병원·요양원, 완도 군청·보건소 등에 실제 방문·영업활동을 했는지 확인하여 주십시오.",
            "2. 차량관리 앱 원본 로그, 블랙박스 원본 파일, 차량 GPS 기록의 생성시각 및 위치정보를 확인하여 주십시오.",
            "3. 법인카드 사용내역 및 급여 지급내역을 확인하여 피고소인이 허위 보고를 통해 취득한 급여 및 비용 상당액을 특정하여 주십시오.",
            "4. 피고소인이 같은 방식으로 2026. 3. 3.부터 2026. 4. 20.까지 다른 근무일에도 허위 보고를 반복하였는지 수사하여 주십시오.",
            "5. 법인카드 식대, 회사 차량, 유류비 등 업무용 자산을 사적 목적으로 사용한 부분에 대하여 업무상배임 또는 업무상횡령 성립 여부도 함께 검토하여 주십시오.",
        ],
    )

    add_heading(doc, "8. 첨부서류", level=1)
    add_numbered(
        doc,
        [
            "1. 김대운 증거 설명서.pdf 1부",
            "2. 차량운행일지 사본 1부",
            "3. 차량관리 앱 운행기록 캡처 1부",
            "4. 차량 블랙박스 캡처 및 원본 영상 저장매체 1부",
            "5. 회사 메신저 캡처 1부",
            "6. 근로계약서 및 업무지시 관련 자료 1부",
            "7. 급여대장 및 급여 이체내역 1부",
            "8. 법인카드 사용내역 및 유류비 내역 1부",
        ],
    )

    doc.add_paragraph()
    date_p = doc.add_paragraph("2026. 5. [  ].")
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_p.runs[0].font.name = "맑은 고딕"
    date_p.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    date_p.runs[0].font.size = Pt(10.5)

    signer = doc.add_paragraph("고소인: [회사명/대표자명]   (서명 또는 인)")
    signer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    signer.runs[0].font.name = "맑은 고딕"
    signer.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    signer.runs[0].font.size = Pt(10.5)

    recipient = doc.add_paragraph("[관할 경찰서장] 귀중")
    recipient.alignment = WD_ALIGN_PARAGRAPH.CENTER
    recipient.runs[0].bold = True
    recipient.runs[0].font.name = "맑은 고딕"
    recipient.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    recipient.runs[0].font.size = Pt(12)

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "제출 전 보강 메모", level=1)
    add_body(
        doc,
        "아래 사항은 고소장 제출 전 빈칸을 채우거나 증빙을 추가하면 수사기관이 피해액과 기망 구조를 더 쉽게 특정할 수 있는 항목입니다.",
    )
    add_numbered(
        doc,
        [
            "1. 고소인 회사명, 대표자, 주소, 사업자등록번호, 담당자 연락처를 입력합니다.",
            "2. 피고소인의 주소, 연락처, 생년월일 또는 주민등록번호 등 인사기록상 정보를 입력합니다.",
            "3. 급여대장과 이체내역으로 2026. 4월분 지급액을 특정하고, 허위·부실 근무 시간 상당액을 산정합니다.",
            "4. 법인카드 사용내역과 영수증으로 2026. 4. 14. 식대 등 업무 관련 비용 지출을 특정합니다.",
            "5. 원본 블랙박스 영상, 차량관리 앱 로그, GPS 기록은 캡처뿐 아니라 원본 파일 또는 관리자 화면 출력본을 함께 제출합니다.",
            "6. ‘잠을 잤다’는 사실은 현재 PDF에서는 직접 장면이 확인되지 않으므로, 블랙박스 영상·목격자 진술 등 직접 증거가 있을 때 별도 기재하는 것이 좋습니다.",
        ],
    )

    doc.save(DOCX_PATH)
    MD_PATH.write_text(complaint_md, encoding="utf-8")
    print(DOCX_PATH)
    print(MD_PATH)


if __name__ == "__main__":
    build_docx()
